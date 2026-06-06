import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOC_ROOT = Path(__file__).resolve().parent
EN_OUT_PATH = DOC_ROOT / "adel_abu_turekei_company_profile_en.docx"
AR_OUT_PATH = DOC_ROOT / "adel_abu_turekei_company_profile_ar.docx"

TABLE_GEOMETRY = DOC_ROOT / Path(
    "C:/Users/ACER NITRO V15/.codex/plugins/cache/openai-primary-runtime/documents/26.601.10930/skills/documents/scripts/table_geometry.py"
)


def load_table_geometry():
    import importlib.util

    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOMETRY)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry()

ACCENT = RGBColor(27, 74, 122)
MUTED = RGBColor(92, 103, 115)
DARK = RGBColor(32, 45, 58)
LIGHT_FILL = "EEF3F7"
TEMP_IMAGE_DIR = Path(tempfile.mkdtemp(prefix="company_profile_imgs_"))

PROJECT_IMAGES = {
    "GlowLabs": DOC_ROOT / "assets/glowlabs/home.png",
    "Nicosia Food": DOC_ROOT / "assets/food/foodHero.png",
    "Oschool": DOC_ROOT / "assets/oschool/schoolHero.png",
    "Magellan": DOC_ROOT / "assets/magellan/heroMagellan.png",
}

EN = {
    "subtitle": "Company Profile",
    "tagline": "Programming, design, marketing, and social media solutions for modern digital products.",
    "definition_title": "Company Definition / Who We Are",
    "definition_body": [
        "ADEL ABU TUREKEI inc is a digital delivery company focused on building modern apps, websites, dashboards, and brand experiences for clients who need practical products ready for launch and growth.",
        "Our work combines product thinking, design execution, and marketing support so each project moves from idea to release with a clear, usable result.",
    ],
    "services_title": "What We Do",
    "services_intro": "We provide end-to-end digital services that cover product planning, interface design, technical delivery, and growth support.",
    "services": [
        "Mobile App Development",
        "Website Development",
        "Dashboard Development",
        "Database Design",
        "UI/UX Design",
        "Branding & Visual Identity",
        "Graphic & Marketing Design",
        "Social Media Management & Advertising",
    ],
    "process_title": "How We Work",
    "process_intro": "Our workflow keeps the client involved from discovery to publishing so scope, feedback, and delivery stay aligned.",
    "process_steps": [
        ("Meet", "Start with a discovery meeting to understand the client needs, identify their goals, and gather the full project requirements."),
        ("Send Proposal", "Send a professional proposal to the client that includes the key highlights, main ideas, project scope, expected deliverables, timeline, and pricing details."),
        ("Get Approval", "Wait for the client to review and consider the proposal, then schedule a final meeting to discuss feedback, confirm the project details, and move forward with the next steps."),
        ("Start Working", "We start working on the project based on the client preferences, turning their goals and vision into a real, practical solution."),
        ("Demo", "Demonstrate the progress achieved so far in stages or batches to ensure the client is satisfied and aligned with the project direction."),
        ("Publish to Store", "Prepare the final release build, complete the listing assets, and publish the product to the stores so the app is ready for distribution."),
    ],
    "projects_title": "Selected Projects",
    "projects_intro": "These projects reflect the range of products we build across marketplaces, food delivery, education, and real estate.",
    "projects": [
        ("GlowLabs", "GlowLabs is a multivendor marketplace application with modern dashboard screens, storefront flows, and commerce management experiences built for production use."),
        ("Nicosia Food", "A food delivery app designed to showcase store catalogs and food listings to customers in Nicosia through a rich, modern, and user-friendly interface."),
        ("Oschool", "Electronic School for digital learning and school management developed using modern web technologies to streamline and enhance traditional school administration and educational work."),
        ("Magellan", "An advanced real estate application designed to list properties and connect buyers with sellers through interactive, social-style chatroom interfaces. The app includes powerful search and filtering features to reduce browsing time and help buyers easily find the most suitable property options."),
    ],
    "pricing_title": "Pricing",
    "pricing_intro": "Project pricing is based on size and delivery complexity. The ranges below are starting guidance for discussion.",
    "pricing_headers": ("Project Size", "Price Range (AED)"),
    "pricing_rows": [
        ("Small project", "1500 - 6000 AED"),
        ("Medium project", "6000 - 15000 AED"),
        ("Large project", "15000 - 30000 AED"),
    ],
    "pricing_note": "All prices are negotiable.",
    "contact_title": "Contact Details",
    "contact_intro": "Tell us about your product, launch timeline, or redesign goals and we will shape the right delivery plan with you.",
    "contact_rows": [
        ("Address", "University St., UAE, Ajman, Jerf"),
        ("Phone", "+971-556-441-299"),
    ],
}

AR = {
    "subtitle": "الملف التعريفي للشركة",
    "tagline": "حلول البرمجة والتصميم والتسويق ووسائل التواصل الاجتماعي للمنتجات الرقمية الحديثة.",
    "definition_title": "تعريف الشركة / من نحن",
    "definition_body": [
        "شركة ADEL ABU TUREKEI inc هي شركة تنفيذ رقمي تركّز على بناء تطبيقات ومواقع ولوحات تحكم وتجارب علامة تجارية حديثة للعملاء الذين يحتاجون إلى منتجات عملية جاهزة للإطلاق والنمو.",
        "يجمع عملنا بين التفكير المنتجـي والتنفيذ التصميمي والدعم التسويقي بحيث ينتقل كل مشروع من الفكرة إلى الإطلاق بنتيجة واضحة وقابلة للاستخدام.",
    ],
    "services_title": "ماذا نقدم",
    "services_intro": "نقدّم خدمات رقمية متكاملة تغطي تخطيط المنتج وتصميم الواجهات والتنفيذ التقني ودعم النمو.",
    "services": [
        "تطوير تطبيقات الجوال",
        "تطوير المواقع الإلكترونية",
        "تطوير لوحات التحكم",
        "تصميم قواعد البيانات",
        "تصميم واجهات وتجربة المستخدم",
        "الهوية البصرية والعلامة التجارية",
        "التصميم الجرافيكي والتسويقي",
        "إدارة وإعلانات وسائل التواصل الاجتماعي",
    ],
    "process_title": "كيف نعمل",
    "process_intro": "تحافظ آلية عملنا على مشاركة العميل من مرحلة الاستكشاف حتى النشر لضمان توافق النطاق والملاحظات والتنفيذ.",
    "process_steps": [
        ("الاجتماع", "نبدأ باجتماع استكشافي لفهم احتياجات العميل وتحديد أهدافه وجمع جميع متطلبات المشروع."),
        ("إرسال العرض", "نرسل عرضًا احترافيًا يتضمن أبرز النقاط والأفكار الرئيسية ونطاق المشروع والمخرجات المتوقعة والجدول الزمني والتسعير."),
        ("الحصول على الموافقة", "ننتظر مراجعة العميل للعرض ثم نعقد اجتماعًا نهائيًا لمناقشة الملاحظات وتأكيد تفاصيل المشروع والانتقال إلى الخطوات التالية."),
        ("بدء العمل", "نبدأ العمل على المشروع وفق تفضيلات العميل، ونحوّل أهدافه ورؤيته إلى حل عملي حقيقي."),
        ("العرض التجريبي", "نعرض التقدم المحقق على مراحل لضمان رضا العميل وتوافقه مع اتجاه المشروع."),
        ("الرفع إلى المتجر", "نجهز النسخة النهائية ونكمل أصول النشر ثم نرفع المنتج إلى المتاجر ليصبح جاهزًا للتوزيع."),
    ],
    "projects_title": "مشاريع مختارة",
    "projects_intro": "تعكس هذه المشاريع نطاق المنتجات التي نبنيها عبر مجالات الأسواق الرقمية وتوصيل الطعام والتعليم والعقارات.",
    "projects": [
        ("GlowLabs", "GlowLabs هو تطبيق سوق متعدد البائعين يضم شاشات لوحة تحكم حديثة وتجارب واجهات متجر وإدارة تجارة مبنية للاستخدام الفعلي."),
        ("Nicosia Food", "تطبيق توصيل طعام صُمم لعرض المتاجر وقوائم الأطعمة للعملاء في نيقوسيا من خلال واجهة غنية وحديثة وسهلة الاستخدام."),
        ("Oschool", "مدرسة إلكترونية للتعلم الرقمي وإدارة المدارس تم تطويرها باستخدام تقنيات ويب حديثة لتبسيط وتحسين الإدارة المدرسية والعمل التعليمي التقليدي."),
        ("Magellan", "تطبيق عقاري متقدم صُمم لعرض العقارات وربط المشترين بالبائعين عبر واجهات محادثة تفاعلية بطابع اجتماعي. يتضمن التطبيق أدوات بحث وتصفية قوية لتقليل وقت التصفح ومساعدة المشترين على العثور بسهولة على أفضل الخيارات المناسبة."),
    ],
    "pricing_title": "الأسعار",
    "pricing_intro": "يعتمد تسعير المشاريع على الحجم وتعقيد التنفيذ، وتمثل الفئات التالية نطاقات استرشادية أولية للنقاش.",
    "pricing_headers": ("حجم المشروع", "نطاق السعر (AED)"),
    "pricing_rows": [
        ("مشروع صغير", "1500 - 6000 AED"),
        ("مشروع متوسط", "6000 - 15000 AED"),
        ("مشروع كبير", "15000 - 30000 AED"),
    ],
    "pricing_note": "جميع الأسعار قابلة للتفاوض.",
    "contact_title": "بيانات التواصل",
    "contact_intro": "أخبرنا عن منتجك أو موعد الإطلاق أو أهداف إعادة التصميم وسنضع معك خطة التنفيذ المناسبة.",
    "contact_rows": [
        ("العنوان", "University St., UAE, Ajman, Jerf"),
        ("الهاتف", "+971-556-441-299"),
    ],
}


def optimized_image(path: Path, max_width=1600):
    image = Image.open(path)
    if image.width <= max_width:
        return str(path)
    ratio = max_width / float(image.width)
    resized = image.resize((max_width, int(image.height * ratio)))
    out_path = TEMP_IMAGE_DIR / f"{path.stem}_optimized.jpg"
    resized.convert("RGB").save(out_path, format="JPEG", quality=88, optimize=True)
    return str(out_path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_bidi(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_rtl(run):
    r_pr = run._r.get_or_add_rPr()
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Arial")


def style_run(run, *, size, bold=False, color=DARK, font="Arial", rtl=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{key}"), font)
    if rtl:
        set_run_rtl(run)


def add_text(paragraph, text, *, size=11, bold=False, color=DARK, rtl=False):
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color, rtl=rtl)
    return run


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "10")
    bottom.set(qn("w:color"), "D0D9E2")
    p_bdr.append(bottom)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in (
        ("Title", 26, 0, 8, ACCENT),
        ("Heading 1", 18, 18, 6, ACCENT),
        ("Heading 2", 14, 14, 4, DARK),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_cover(doc, data, rtl=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_bidi(p)
    add_text(p, "ADEL ABU TUREKEI inc", size=24, bold=True, color=ACCENT, rtl=rtl)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_bidi(p)
    add_text(p, data["subtitle"], size=12, bold=True, color=MUTED, rtl=rtl)
    p.paragraph_format.space_after = Pt(14)

    hero = doc.add_picture(optimized_image(PROJECT_IMAGES["GlowLabs"]), width=Inches(6.1))
    hero.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_bidi(p)
    add_text(p, data["tagline"], size=11, color=DARK, rtl=rtl)
    p.paragraph_format.space_after = Pt(10)

    divider = doc.add_paragraph()
    add_rule(divider)
    divider.paragraph_format.space_after = Pt(8)


def add_heading(doc, text, rtl=False):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_bidi(p)
    add_text(p, text, size=18 if not rtl else 16, bold=True, color=ACCENT, rtl=rtl)


def add_paragraph_text(doc, text, rtl=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_bidi(p)
    add_text(p, text, size=11, color=DARK, rtl=rtl)


def add_bullets(doc, items, rtl=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, item, size=11, color=DARK, rtl=rtl)


def add_process_table(doc, data, rtl=False):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = data["pricing_headers"] if False else (
        ("الخطوة", "الشرح") if rtl else ("Step", "Details")
    )
    for idx, label in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p)
        add_text(p, label, size=11, bold=True, color=ACCENT, rtl=rtl)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)

    for title, body in data["process_steps"]:
        row = table.add_row().cells
        p1 = row[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p1)
        add_text(p1, title, size=10.5, bold=True, color=DARK, rtl=rtl)
        p2 = row[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p2)
        add_text(p2, body, size=10.5, color=DARK, rtl=rtl)

    if rtl:
        table_geometry.apply_table_geometry(table, [2500, 6860], table_width_dxa=9360, indent_dxa=120)
    else:
        table_geometry.apply_table_geometry(table, [2200, 7160], table_width_dxa=9360, indent_dxa=120)


def add_projects(doc, data, rtl=False):
    for name, body in data["projects"]:
        p = doc.add_paragraph(style="Heading 2")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p)
        add_text(p, name, size=14, bold=True, color=DARK, rtl=rtl)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        if rtl:
            text_cell, image_cell = table.rows[0].cells
        else:
            image_cell, text_cell = table.rows[0].cells

        image_p = image_cell.paragraphs[0]
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.add_run().add_picture(optimized_image(PROJECT_IMAGES[name]), width=Inches(2.4))

        text_p = text_cell.paragraphs[0]
        text_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(text_p)
        add_text(text_p, body, size=10.5, color=DARK, rtl=rtl)
        table_geometry.apply_table_geometry(table, [6560, 2800] if rtl else [2800, 6560], table_width_dxa=9360, indent_dxa=120)
        doc.add_paragraph()


def add_pricing(doc, data, rtl=False):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = data["pricing_headers"]
    for idx, label in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p)
        add_text(p, label, size=11, bold=True, color=ACCENT, rtl=rtl)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)

    for left, right in data["pricing_rows"]:
        row = table.add_row().cells
        p1 = row[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_bidi(p1)
        add_text(p1, left, size=11, color=DARK, rtl=rtl)
        p2 = row[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p2, right, size=11, bold=True, color=DARK, rtl=False)
    table_geometry.apply_table_geometry(table, [4320, 5040], table_width_dxa=9360, indent_dxa=120)
    add_paragraph_text(doc, data["pricing_note"], rtl=rtl)


def add_contact(doc, data, rtl=False):
    table = doc.add_table(rows=len(data["contact_rows"]), cols=2)
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, data["contact_rows"]):
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            if rtl and idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_bidi(p)
                add_text(p, text, size=11, bold=True if idx == 0 else False, color=ACCENT if idx == 0 else DARK, rtl=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if not rtl or idx == 1 else WD_ALIGN_PARAGRAPH.RIGHT
                add_text(p, text, size=11, bold=True if idx == 0 else False, color=ACCENT if idx == 0 else DARK, rtl=False)
    table_geometry.apply_table_geometry(table, [2800, 6560], table_width_dxa=9360, indent_dxa=120)


def build_profile(data, output_path, rtl=False):
    doc = Document()
    configure_document(doc)
    add_cover(doc, data, rtl=rtl)

    add_heading(doc, data["definition_title"], rtl=rtl)
    for body in data["definition_body"]:
        add_paragraph_text(doc, body, rtl=rtl)

    add_heading(doc, data["services_title"], rtl=rtl)
    add_paragraph_text(doc, data["services_intro"], rtl=rtl)
    add_bullets(doc, data["services"], rtl=rtl)

    add_heading(doc, data["process_title"], rtl=rtl)
    add_paragraph_text(doc, data["process_intro"], rtl=rtl)
    add_process_table(doc, data, rtl=rtl)

    add_heading(doc, data["projects_title"], rtl=rtl)
    add_paragraph_text(doc, data["projects_intro"], rtl=rtl)
    add_projects(doc, data, rtl=rtl)

    add_heading(doc, data["pricing_title"], rtl=rtl)
    add_paragraph_text(doc, data["pricing_intro"], rtl=rtl)
    add_pricing(doc, data, rtl=rtl)

    add_heading(doc, data["contact_title"], rtl=rtl)
    add_paragraph_text(doc, data["contact_intro"], rtl=rtl)
    add_contact(doc, data, rtl=rtl)

    doc.save(output_path)


def main():
    build_profile(EN, EN_OUT_PATH, rtl=False)
    build_profile(AR, AR_OUT_PATH, rtl=True)


if __name__ == "__main__":
    main()
